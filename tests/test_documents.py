import sys
from pathlib import Path
import shutil
import subprocess

from docx import Document
from PIL import Image
from pypdf import PdfReader, PdfWriter
import pytest

PROJECT_ROOT = Path(__file__).parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

import grader_core.documents as documents
from grader_core.documents import (
    AMBIGUOUS_ANSWER_FILES,
    DOCUMENT_CONVERSION_FAILED,
    DOCUMENT_CONVERSION_UNAVAILABLE,
    MISSING_ANSWER_FILE,
    NormalizationFailure,
    NormalizedDocument,
    PageContent,
    SubmissionFiles,
    chunk_visual_pages,
    discover_submissions,
    extract_html_blocks,
    normalize_document,
    render_visual_pages,
    select_visual_pages,
)


def _write(path: Path, content: str = "sample") -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return path


def test_discover_submissions_excludes_question_and_declaration_files(
    tmp_path: Path,
) -> None:
    assignment_root = tmp_path / "StudentAnswer_ISYS6599"
    student_folder = assignment_root / "2902737810_LUK SEKAR DADARI"
    answer_path = _write(student_folder / "answer.pdf")
    _write(student_folder / "Question.pdf")
    _write(student_folder / "ai_usage_declaration.pdf")
    question_path = _write(student_folder / "Question.html", "<p>Question text</p>")

    submissions = discover_submissions([assignment_root])

    assert submissions == [
        SubmissionFiles(
            student_id="2902737810",
            student_name="LUK SEKAR DADARI",
            assignment_root=assignment_root,
            folder=student_folder,
            answer_path=answer_path,
            question_path=question_path,
            review_reasons=(),
        )
    ]


def test_discover_submissions_flags_multiple_answer_candidates(tmp_path: Path) -> None:
    assignment_root = tmp_path / "StudentAnswer_ISYS6599"
    student_folder = assignment_root / "2902737810_LUK SEKAR DADARI"
    _write(student_folder / "answer.pdf")
    _write(student_folder / "revised_answer.docx")

    submission = discover_submissions([assignment_root])[0]

    assert submission.answer_path is None
    assert submission.review_reasons == (AMBIGUOUS_ANSWER_FILES,)


def test_discover_submissions_flags_missing_answer_file(tmp_path: Path) -> None:
    assignment_root = tmp_path / "StudentAnswer_ISYS6599"
    student_folder = assignment_root / "2902737810_LUK SEKAR DADARI"
    _write(student_folder / "Question.pdf")

    submission = discover_submissions([assignment_root])[0]

    assert submission.answer_path is None
    assert submission.review_reasons == (MISSING_ANSWER_FILE,)


def test_discover_submissions_uses_assignment_root_question_fallback(
    tmp_path: Path,
) -> None:
    assignment_root = tmp_path / "StudentAnswer_ISYS6599"
    student_folder = assignment_root / "2902737810_LUK SEKAR DADARI"
    answer_path = _write(student_folder / "answer.pdf")
    question_path = _write(
        assignment_root / "Question.html", "<p>Weekly question</p>"
    )

    submission = discover_submissions([assignment_root])[0]

    assert submission.assignment_root == assignment_root
    assert submission.question_path == question_path
    assert submission.answer_path == answer_path


def test_discover_submissions_parses_group_folder_identity(tmp_path: Path) -> None:
    assignment_root = tmp_path / "StudentAnswer_ISYS6310035"
    group_folder = assignment_root / "Group_Group-1"
    answer_path = _write(group_folder / "submission.pdf")

    submission = discover_submissions([assignment_root])[0]

    assert submission.student_id == "group-1"
    assert submission.student_name == "Group 1"
    assert submission.answer_path == answer_path


def test_discover_submissions_rejects_non_assignment_roots(tmp_path: Path) -> None:
    non_assignment_root = tmp_path / "uploads"
    non_assignment_root.mkdir()

    submissions = discover_submissions([non_assignment_root])

    assert submissions == []


def test_extract_html_blocks_preserves_numbered_questions() -> None:
    html = """
    <h2>Question 1</h2><p>Explain COBIT.</p>
    <ol><li>Define the framework.</li><li>Give an example.</li></ol>
    <table><tr><th>Criterion</th><td>Evidence</td></tr></table>
    <p>Question 2<br>Apply ISO 31000.</p>
    """

    assert extract_html_blocks(html) == (
        "Question 1\n\n"
        "Explain COBIT.\n\n"
        "Define the framework.\n"
        "Give an example.\n\n"
        "Criterion\n"
        "Evidence\n\n"
        "Question 2\n"
        "Apply ISO 31000."
    )


def test_normalize_pdf_preserves_page_count_and_page_diagnostics(
    tmp_path: Path, synthetic_pdf_files: dict[str, Path]
) -> None:
    source_path = tmp_path / "combined.pdf"
    writer = PdfWriter()
    writer.append(PdfReader(synthetic_pdf_files["text"]))
    writer.append(PdfReader(synthetic_pdf_files["table"]))
    writer.append(PdfReader(synthetic_pdf_files["image_only"]))
    with source_path.open("wb") as stream:
        writer.write(stream)

    normalized = normalize_document(source_path, tmp_path / "run-temp")

    assert normalized.source_path == source_path
    assert normalized.pdf_path == source_path
    assert normalized.source_sha256 == normalized.pdf_sha256
    assert [page.number for page in normalized.pages] == [1, 2, 3]
    assert "Jawaban normal" in normalized.pages[0].text
    assert normalized.pages[1].has_table is True
    assert normalized.pages[2].text == ""
    assert len(normalized.pages[2].image_hashes) == 1


def test_normalize_docx_runs_bounded_libreoffice_conversion(
    tmp_path: Path,
    synthetic_pdf_files: dict[str, Path],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source_path = tmp_path / "answer.docx"
    document = Document()
    document.add_paragraph("Answer content")
    document.save(source_path)
    temporary_dir = tmp_path / "run-temp"
    captured: dict[str, object] = {}

    def fake_run(command: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        captured["command"] = command
        captured["kwargs"] = kwargs
        output_path = temporary_dir / "answer.pdf"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(synthetic_pdf_files["text"], output_path)
        return subprocess.CompletedProcess(command, 0, "", "")

    monkeypatch.setattr(documents.shutil, "which", lambda _name: "libreoffice")
    monkeypatch.setattr(documents.subprocess, "run", fake_run)

    normalized = normalize_document(source_path, temporary_dir)

    assert normalized.pdf_path == temporary_dir / "answer.pdf"
    assert normalized.pdf_path != source_path
    assert captured["command"] == [
        "libreoffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(temporary_dir),
        str(source_path),
    ]
    assert captured["kwargs"] == {
        "capture_output": True,
        "check": False,
        "text": True,
        "timeout": 60,
    }


def test_normalize_docx_flags_unavailable_libreoffice(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source_path = tmp_path / "answer.docx"
    source_path.write_bytes(b"not-converted")
    monkeypatch.setattr(documents.shutil, "which", lambda _name: None)

    normalized = normalize_document(source_path, tmp_path / "run-temp")

    assert normalized == NormalizationFailure(
        source_path=source_path,
        review_reason=DOCUMENT_CONVERSION_UNAVAILABLE,
    )


def test_normalize_docx_flags_failed_conversion(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source_path = tmp_path / "answer.docx"
    source_path.write_bytes(b"not-converted")
    monkeypatch.setattr(documents.shutil, "which", lambda _name: "libreoffice")
    monkeypatch.setattr(
        documents.subprocess,
        "run",
        lambda command, **_kwargs: subprocess.CompletedProcess(command, 1, "", "failed"),
    )

    normalized = normalize_document(source_path, tmp_path / "run-temp")

    assert normalized == NormalizationFailure(
        source_path=source_path,
        review_reason=DOCUMENT_CONVERSION_FAILED,
    )


def test_repeated_template_image_does_not_select_every_text_page(
    tmp_path: Path,
) -> None:
    document = NormalizedDocument(
        source_path=tmp_path / "answer.pdf",
        pdf_path=tmp_path / "answer.pdf",
        source_sha256="source",
        pdf_sha256="pdf",
        pages=tuple(
            PageContent(
                number=page_number,
                text="A sufficiently complete answer " * 10,
                image_hashes=("repeated-logo",),
                has_table=False,
                has_drawings=False,
            )
            for page_number in range(1, 4)
        ),
    )

    assert select_visual_pages(document, min_text_chars=100) == ()


def test_unique_diagram_and_image_only_pages_are_selected(tmp_path: Path) -> None:
    document = NormalizedDocument(
        source_path=tmp_path / "answer.pdf",
        pdf_path=tmp_path / "answer.pdf",
        source_sha256="source",
        pdf_sha256="pdf",
        pages=(
            PageContent(1, "Complete answer " * 20, ("logo",), False, False),
            PageContent(2, "Complete answer " * 20, ("logo",), False, False),
            PageContent(3, "Complete answer " * 20, ("diagram",), False, False),
            PageContent(4, "", ("answer-image",), False, False),
        ),
    )

    assert select_visual_pages(document, min_text_chars=100) == (3, 4)


def test_table_and_excessive_drawing_pages_are_selected(tmp_path: Path) -> None:
    document = NormalizedDocument(
        source_path=tmp_path / "answer.pdf",
        pdf_path=tmp_path / "answer.pdf",
        source_sha256="source",
        pdf_sha256="pdf",
        pages=(
            PageContent(1, "Complete answer " * 20, (), True, False),
            PageContent(2, "Complete answer " * 20, (), False, True, 11),
            PageContent(3, "Complete answer " * 20, (), False, True, 2),
        ),
    )

    assert select_visual_pages(document, min_text_chars=100, max_drawing_count=10) == (
        1,
        2,
    )


def test_visual_pages_are_chunked_by_four() -> None:
    assert chunk_visual_pages(range(1, 10)) == (
        (1, 2, 3, 4),
        (5, 6, 7, 8),
        (9,),
    )


def test_render_visual_pages_writes_bounded_pngs(
    tmp_path: Path, synthetic_pdf_files: dict[str, Path]
) -> None:
    rendered = render_visual_pages(
        synthetic_pdf_files["text"],
        (1,),
        tmp_path / "rendered",
        scale=2.0,
        max_pixels=50_000,
    )

    assert len(rendered) == 1
    assert rendered[0].page_number == 1
    assert rendered[0].path.name == "page-0001.png"
    assert rendered[0].path.is_file()
    with Image.open(rendered[0].path) as image:
        assert image.format == "PNG"
        assert image.width * image.height <= 50_000
        assert image.width > 0
        assert image.height > 0
